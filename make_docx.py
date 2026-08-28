"""
make_docx.py -- render FINAL_REPORT.md to a Word document.

Produces CS175_Final_Report_Group3.docx with real heading styles, real tables,
monospaced pseudocode, superscript exponents, and the figures already placed and
sized. Open it in Google Docs (File > Open, or drop it in Drive) and it converts
cleanly, so there is no retyping step.

    pip3 install python-docx
    python3 make_docx.py
"""
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(REPO, "FINAL_REPORT.md")
OUT = os.path.join(REPO, "CS175_Final_Report_Group3.docx")

FIG_WIDTH = Inches(4.9)          # about 75% of the 6.5 inch text column
BODY_PT = 11
MONO_PT = 9

doc = Document()

# page setup and base styles
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1)

normal = doc.styles["Normal"]
normal.font.name = "Georgia"
normal.font.size = Pt(BODY_PT)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for name, size in (("Heading 1", 16), ("Heading 2", 12.5)):
    st = doc.styles[name]
    st.font.name = "Arial"
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    st.paragraph_format.space_before = Pt(14 if size > 14 else 11)
    st.paragraph_format.space_after = Pt(4)


INLINE = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*)')

def add_runs(par, text):
    """Emit text, honouring `code` and **bold**, and superscripting 9^64 style."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1])
            r.font.name = "Courier New"
            r.font.size = Pt(BODY_PT - 1.5)
        elif chunk.startswith("**") and chunk.endswith("**"):
            par.add_run(chunk[2:-2]).bold = True
        else:
            # split out exponents so they can be raised
            for piece in re.split(r'(\d+\^\d+)', chunk):
                m = re.fullmatch(r'(\d+)\^(\d+)', piece or "")
                if m:
                    par.add_run(m.group(1))
                    sup = par.add_run(m.group(2))
                    sup.font.superscript = True
                elif piece:
                    par.add_run(piece)


def add_table(rows):
    header, body = rows[0], rows[2:]      # rows[1] is the |---| separator
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, cell in enumerate(t.rows[0].cells):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(header[i])
        r.bold = True
        r.font.size = Pt(BODY_PT - 1.5)
        r.font.name = "Arial"
    for row in body:
        cells = t.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(BODY_PT - 1.5)
            r.font.name = "Arial"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code(lines):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(8)
    pf.left_indent = Inches(0.25)
    pf.line_spacing = 1.0
    for i, ln in enumerate(lines):
        r = p.add_run(("" if i == 0 else "\n") + ln)
        r.font.name = "Courier New"
        r.font.size = Pt(MONO_PT)


# ---- walk the markdown ----------------------------------------------------- #
lines = open(SRC, encoding="utf-8").read().splitlines()
# everything before the first "## " heading is the cover block, where each line
# stands alone (one team member per line) rather than wrapping into a paragraph
first_heading = next((k for k, l in enumerate(lines) if l.startswith("## ")), len(lines))
i = 0
figno = 0
while i < len(lines):
    ln = lines[i]
    stripped = ln.strip()

    if not stripped:
        i += 1
        continue

    if stripped.startswith("<!--"):
        i += 1
        continue

    # title block: use the built-in Title style so Google Docs imports it as a
    # real title rather than a bold paragraph it may reflow or truncate
    if stripped.startswith("# "):
        p = doc.add_paragraph(style="Title")
        # the Title style draws a bottom border; remove it
        pPr = p._p.get_or_add_pPr()
        for b in pPr.findall(qn('w:pBdr')):
            pPr.remove(b)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(stripped[2:])
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(19)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        i += 1
        continue

    if stripped.startswith("### "):
        doc.add_paragraph(stripped[4:], style="Heading 2")
        i += 1
        continue

    if stripped.startswith("## "):
        doc.add_paragraph(stripped[3:], style="Heading 1")
        i += 1
        continue

    # figure placeholder -> insert the image itself
    m = re.match(r'\[\[INSERT FIGURE (\d+) HERE:\s*([^\]]+)\]\]', stripped)
    if m:
        path = os.path.join(REPO, m.group(2).strip())
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        if os.path.exists(path):
            p.add_run().add_picture(path, width=FIG_WIDTH)
            figno += 1
        else:
            p.add_run(f"[missing image: {path}]")
        i += 1
        continue

    # table
    if stripped.startswith("|"):
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
            i += 1
        add_table(rows)
        continue

    # indented block: pseudocode or a display equation
    if ln.startswith("    "):
        block = []
        while i < len(lines) and (lines[i].startswith("    ") or not lines[i].strip()):
            if not lines[i].strip():
                if i + 1 < len(lines) and lines[i + 1].startswith("    "):
                    block.append("")
                    i += 1
                    continue
                break
            block.append(lines[i][4:])
            i += 1
        add_code(block)
        continue

    # bullet list
    if stripped.startswith("- "):
        while i < len(lines) and lines[i].strip().startswith("- "):
            text = lines[i].strip()[2:]
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip() \
                    and not lines[i].strip().startswith("- "):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, text)
        continue

    # numbered list
    if re.match(r'^\d+\.\s', stripped):
        while i < len(lines) and re.match(r'^\d+\.\s', lines[i].strip()):
            text = re.sub(r'^\d+\.\s+', '', lines[i].strip())
            i += 1
            while i < len(lines) and lines[i].startswith("   ") and lines[i].strip() \
                    and not re.match(r'^\d+\.\s', lines[i].strip()):
                text += " " + lines[i].strip()
                i += 1
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, text)
        continue

    # cover block: one paragraph per line, no joining
    if i < first_heading:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, stripped)
        i += 1
        continue

    # ordinary paragraph, joining its wrapped lines
    para = [stripped]
    i += 1
    while i < len(lines) and lines[i].strip() and not lines[i].startswith("    ") \
            and not lines[i].strip().startswith(("#", "|", "- ", "[[", "<!--")) \
            and not re.match(r'^\d+\.\s', lines[i].strip()):
        para.append(lines[i].strip())
        i += 1
    text = " ".join(para)
    p = doc.add_paragraph()
    # figure captions read better small and italic
    if re.match(r'^Figure \d+\.', text):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(BODY_PT - 1.5)
    else:
        add_runs(p, text)

doc.save(OUT)
print(f"wrote {OUT}")
print(f"  {figno} figures embedded at {FIG_WIDTH.inches:.2f} inches wide")
